"""Word 문서(.doc / .docx) 로더.

지원 파일:
  *.docx → python-docx 사용 (추가 설치 없이 동작)
  *.doc  → win32com.client (Windows + Microsoft Word 필요)
           없으면 LibreOffice soffice CLI로 fallback

청킹 전략:
  '제N조' 패턴 기준으로 조(article) 단위 청킹.
  조 번호 이전 서문/제목은 별도 intro 청크로 저장.
"""

import hashlib
import re
import subprocess
import tempfile
from pathlib import Path

from langchain_core.documents import Document
from loguru import logger

# 줄 시작에서만 조항 경계로 인식 — 본문 내 교차참조("제3조에 따른...")는 무시
# 패턴: 줄 시작 + 제N조 또는 제N조의N + 선택적 공백/괄호 제목
ARTICLE_PATTERN = re.compile(r"(?m)^(제\s*\d+\s*조(?:의\d+)?(?:\s*\([^)]+\))?)")


def _normalize(text: str) -> str:
    """Windows CRLF → LF, 연속 공백 정리."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _chunk_by_article(text: str, source_name: str) -> list[Document]:
    """줄 시작 기준 조(제N조) 단위로 청킹.

    본문 내 교차참조("제3조에 따른...")는 분리 기준에서 제외된다.
    조 번호 이전 서문/제목은 'intro' 청크로 따로 저장한다.
    """
    text = _normalize(text)
    parts = ARTICLE_PATTERN.split(text)
    # split 결과: [intro, "제1조", "내용", "제2조", "내용", ...]
    # parts[0] = 조 번호 이전 텍스트 (intro)
    # parts[1::2] = 조 번호들, parts[2::2] = 각 조 본문

    documents: list[Document] = []

    # intro 처리 (조 번호 이전 서문/제목)
    intro = parts[0].strip()
    if intro:
        chunk_id = hashlib.md5(f"{source_name}:intro".encode()).hexdigest()
        documents.append(Document(
            page_content=intro,
            metadata={
                "title": source_name,
                "article": "서문",
                "chunk_id": chunk_id,
                "source_type": "law_statute",
            },
        ))

    # 조 단위 처리
    article_labels = parts[1::2]
    article_bodies = parts[2::2]

    for label, body in zip(article_labels, article_bodies):
        label = label.strip()
        content = (label + "\n" + body.strip()).strip()
        if not content:
            continue
        chunk_id = hashlib.md5(f"{source_name}:{label}".encode()).hexdigest()
        documents.append(Document(
            page_content=content,
            metadata={
                "title": source_name,
                "article": label,
                "chunk_id": chunk_id,
                "source_type": "law_statute",
            },
        ))

    return documents


# ──────────────────────────────────────────────────
# .docx 로더 (python-docx)
# ──────────────────────────────────────────────────

def _load_docx(file_path: Path) -> str:
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(str(file_path))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as e:
        logger.error(f"[DOC로더] docx 읽기 실패 {file_path.name}: {e}")
        return ""


# ──────────────────────────────────────────────────
# .doc 로더 — win32com (Windows + Word 필요)
# ──────────────────────────────────────────────────

def _load_doc_win32(file_path: Path) -> str:
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(file_path.resolve()))
        text = doc.Content.Text
        doc.Close(False)
        word.Quit()
        return text
    except ImportError:
        return ""
    except Exception as e:
        logger.warning(f"[DOC로더] win32com 실패 {file_path.name}: {e}")
        return ""


# ──────────────────────────────────────────────────
# .doc 로더 — LibreOffice soffice CLI (fallback)
# ──────────────────────────────────────────────────

def _load_doc_soffice(file_path: Path) -> str:
    """LibreOffice로 .doc → .txt 변환 후 읽기."""
    try:
        with tempfile.TemporaryDirectory() as tmp:
            result = subprocess.run(
                ["soffice", "--headless", "--convert-to", "txt:Text",
                 "--outdir", tmp, str(file_path.resolve())],
                capture_output=True, timeout=30,
            )
            if result.returncode != 0:
                return ""
            txt_file = Path(tmp) / (file_path.stem + ".txt")
            if txt_file.exists():
                return txt_file.read_text(encoding="utf-8", errors="ignore")
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    return ""


# ──────────────────────────────────────────────────
# 공개 API
# ──────────────────────────────────────────────────

def load_doc_file(file_path: str | Path) -> list[Document]:
    """단일 .doc / .docx 파일을 Document 리스트로 변환.

    처리 우선순위:
      .docx → python-docx
      .doc  → win32com → soffice → 경고 후 빈 리스트
    """
    file_path = Path(file_path)
    if not file_path.exists():
        logger.warning(f"[DOC로더] 파일 없음: {file_path}")
        return []

    ext = file_path.suffix.lower()
    text = ""

    if ext == ".docx":
        text = _load_docx(file_path)
    elif ext == ".doc":
        text = _load_doc_win32(file_path)
        if not text:
            logger.info(f"[DOC로더] win32com 미작동, LibreOffice 시도: {file_path.name}")
            text = _load_doc_soffice(file_path)
        if not text:
            logger.warning(
                f"[DOC로더] {file_path.name} 로드 실패. "
                "Microsoft Word 또는 LibreOffice가 설치되어 있어야 합니다."
            )
            return []
    else:
        logger.warning(f"[DOC로더] 미지원 확장자: {ext}")
        return []

    if not text.strip():
        logger.warning(f"[DOC로더] 빈 문서: {file_path.name}")
        return []

    source_name = file_path.stem
    chunks = _chunk_by_article(text, source_name)
    logger.info(f"[DOC로더] {file_path.name} → {len(chunks)}개 청크")
    return chunks


def load_doc_dir(folder_path: str | Path) -> list[Document]:
    """폴더 내 모든 .doc / .docx 파일 로드.

    Args:
        folder_path: 법률 폴더 경로.

    Returns:
        Document 리스트.
    """
    folder_path = Path(folder_path)
    doc_files = sorted(
        list(folder_path.glob("*.doc")) + list(folder_path.glob("*.docx"))
    )

    if not doc_files:
        logger.warning(f"[DOC로더] .doc/.docx 파일 없음: {folder_path}")
        return []

    logger.info(f"[DOC로더] {len(doc_files)}개 파일 로드 시작: {folder_path.name}")

    documents: list[Document] = []
    for df in doc_files:
        documents.extend(load_doc_file(df))

    logger.info(f"[DOC로더] {folder_path.name} → 총 {len(documents)}개 Document")
    return documents
